from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

print('Criando arquivo de referencia Word para formatacao ABNT...')

# Criar documento Word
doc = Document()

# Configurar estilos ABNT
def configurar_estilos_abnt(doc):
    # Estilo Normal - Times New Roman 12pt, espaçamento 1.5
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.first_line_indent = Inches(0.49)  # 1.25cm = 0.49 inches
    
    # Título 1 - Seções principais
    try:
        style_h1 = doc.styles['Heading 1']
    except:
        style_h1 = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    
    font_h1 = style_h1.font
    font_h1.name = 'Times New Roman'
    font_h1.size = Pt(12)
    font_h1.bold = True
    
    para_h1 = style_h1.paragraph_format
    para_h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_h1.space_before = Pt(18)
    para_h1.space_after = Pt(12)
    para_h1.keep_with_next = True
    
    # Título 2 - Subseções
    try:
        style_h2 = doc.styles['Heading 2']
    except:
        style_h2 = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    
    font_h2 = style_h2.font
    font_h2.name = 'Times New Roman'
    font_h2.size = Pt(12)
    font_h2.bold = True
    
    para_h2 = style_h2.paragraph_format
    para_h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_h2.space_before = Pt(12)
    para_h2.space_after = Pt(6)
    para_h2.keep_with_next = True
    
    # Legenda de figura/tabela
    try:
        style_caption = doc.styles['Caption']
    except:
        style_caption = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    
    font_caption = style_caption.font
    font_caption.name = 'Times New Roman'
    font_caption.size = Pt(10)
    
    para_caption = style_caption.paragraph_format
    para_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_caption.space_before = Pt(6)
    para_caption.space_after = Pt(6)
    para_caption.keep_with_next = False

# Aplicar estilos
configurar_estilos_abnt(doc)

# Adicionar conteúdo exemplo para definir estilos
doc.add_heading('1 INTRODUÇÃO', level=1)
p1 = doc.add_paragraph('Este é um exemplo de parágrafo normal com formatação ABNT. O texto deve estar em Times New Roman 12pt, com espaçamento 1,5 linhas e recuo de primeira linha de 1,25 cm.')

doc.add_heading('1.1 OBJETIVOS', level=2)
p2 = doc.add_paragraph('Este é outro parágrafo exemplo para demonstrar a formatação correta.')

doc.add_heading('2 METODOLOGIA', level=1)
p3 = doc.add_paragraph('Exemplo de texto da metodologia.')

# Configurar margens do documento (ABNT: 3cm superior e esquerda, 2cm inferior e direita)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.181)    # 3cm exato
    section.left_margin = Inches(1.181)   # 3cm exato  
    section.bottom_margin = Inches(0.787) # 2cm exato
    section.right_margin = Inches(0.787)  # 2cm exato

# Salvar arquivo de referência
reference_path = r'C:\Users\BrunoGaspariniBaller\OneDrive - HAND\Documentos\TCC\Overleaf\TCC_RiskParity\docs\Overleaf\reference.docx'
doc.save(reference_path)

print(f'Arquivo de referencia criado: {reference_path}')
print('Configuracoes ABNT aplicadas:')
print('- Fonte: Times New Roman 12pt')
print('- Espacamento: 1.5 linhas')
print('- Margens: 3cm (sup/esq), 2cm (inf/dir)')  
print('- Recuo primeira linha: 1.25cm')
print('- Titulos numerados e formatados')